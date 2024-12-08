import os
import re
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor  # 用于设置字号
from docx.enum.text import WD_ALIGN_PARAGRAPH  # 用于设置段落格式
from collections import defaultdict

folder_path = "/Users/yuki/OneDrive/Workspace/课程文件/环境法/环境资源法命题收集"  # 需要指定文件夹路径

# 获取文件夹内所有 docx 文件
docx_files = [f for f in os.listdir(folder_path) if f.endswith(".docx")]


def extract_chapter_number(filename):
    # 用正则表达式提取文件名中的数字部分，假设数字是由点分隔的
    match = re.match(r"(\d+(\.\d+)+)", filename)
    if match:
        return [
            tuple(map(int, match.group(1).split(".")))
        ]  # 将章节编号转为元组形式以便于比较
    return []


# 按照章节编号排序文件
sorted_files = sorted(docx_files, key=extract_chapter_number)
print(sorted_files)

# 读取所有文件内容
documents = [Document(os.path.join(folder_path, file)) for file in sorted_files]

# 提取每个文档的内容（仅提取文本部分）
doc_texts = []
for doc in documents:
    doc_text = []
    for para in doc.paragraphs:
        doc_text.append(para.text.strip())
    doc_texts.append(doc_text)


# 创建一个字典来存储按题型分类的题目
question_by_type = defaultdict(list)

# 题型标签列表
question_types = ["单选题", "多选题", "判断题"]
question_types_labals = {
    "单选题": "[单选题][2分][难度3]",
    "多选题": "[多选题][2分][难度3]",
    "判断题": "[判断题][1分][难度1]",
}
question_types_titles = {
    "单选题": "一、单选题（一共",
    "多选题": "三、多选题（一共",
    "判断题": "二、判断题（一共",
}
question_count = {
    "单选题": 0,
    "多选题": 0,
    "判断题": 0,
}

# 解析每个文档的内容，分类题目
for doc_text in doc_texts:
    current_type = None
    for line in doc_text:
        if line.strip():
            line = line.strip()
            # 判断是否为题型标签
            if any(question_type in line for question_type in question_types):
                current_type = next(
                    question_type
                    for question_type in question_types
                    if question_type in line
                )

                if question_count[current_type]:
                    question_by_type[current_type].append("")

                question_by_type[current_type].append(
                    question_types_labals[current_type]
                )
                question_count[current_type] += 1
            elif current_type:
                # 将题目按题型添加到字典中
                question_by_type[current_type].append(line)

# # 查看分类结果的一部分
# print(dict(list(question_by_type.items())[:1]))  # 查看前两个题型的内容

# 创建新的 Word 文档
new_doc = Document()

# 获取 "Normal" 样式
style = new_doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(10.5)
font.color.rgb = RGBColor(0, 0, 0)

# 定义题型的排序优先级
question_type_order = {
    "单选题": 1,
    "判断题": 2,
    "多选题": 3,
}

# 按题型顺序排列
sorted_question_types = sorted(
    question_by_type.keys(), key=lambda x: question_type_order.get(x, 4)
)

paragraph = new_doc.add_paragraph(
    "文档四：期末考试题库（按照单选、判断、多选的题型排列，按照章节顺序）"
)
run = paragraph.runs[0]
run.font.color.rgb = RGBColor(255, 0, 0)
paragraph_format = paragraph.paragraph_format
paragraph_format.line_spacing = 1  # 行距为1
paragraph_format.space_before = 0  # 段前距为0
paragraph_format.space_after = 0  # 段后距为0

# 按题型分类添加题目
for question_type, questions in question_by_type.items():
    # 添加题型标题
    paragraph = new_doc.add_paragraph(
        f"{question_types_titles[question_type]}{question_count[question_type]}道）"
    )
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = 1  # 行距为1
    paragraph_format.space_before = 0  # 段前距为0
    paragraph_format.space_after = 0  # 段后距为0

    # 添加每个题目
    for question in questions:
        # 添加段落
        paragraph = new_doc.add_paragraph(question)
        # if question.strip():
        #     # 设置字体为宋体，字号为五号，行距为1，段前段后缩进为0
        #     run = paragraph.runs[0]
        #     run.font.name = "Calibri"
        #     run._r.rPr.rFonts.set("all", "宋体")  # 确保字体设置为宋体（针对中文）
        #     run.font.size = Pt(10.5)  # 设置字号为五号（10.5 pt）

        # 设置段落格式
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = 1  # 行距为1
        paragraph_format.space_before = 0  # 段前距为0
        paragraph_format.space_after = 0  # 段后距为0
        paragraph_format.left_indent = 0  # 左缩进为0
        paragraph_format.right_indent = 0  # 右缩进为0


# 保存新的文档
output_path = "./output.docx"
new_doc.save(output_path)
